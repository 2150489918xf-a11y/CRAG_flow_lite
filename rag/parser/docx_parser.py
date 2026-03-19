"""
DOCX 解析器
"""
import logging
import re

logger = logging.getLogger(__name__)


def parse(filename, binary=None):
    """
    解析 DOCX 文件，返回 (sections, tables)
    sections = [(text, tag), ...]
    tables = [html_table_str, ...]
    """
    from docx import Document
    import io

    try:
        if binary:
            doc = Document(io.BytesIO(binary))
        else:
            doc = Document(filename)
    except Exception as e:
        logger.error(f"Failed to open DOCX {filename}: {e}")
        return [], []

    sections = []
    tables = []

    # 提取段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        tag = "text"
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if "heading" in style_name or "title" in style_name:
                tag = "title"
        sections.append((text, tag))

    # 提取表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            header = rows[0]
            html = "<table>\n<tr>" + "".join(f"<th>{h}</th>" for h in header) + "</tr>\n"
            for row in rows[1:]:
                html += "<tr>" + "".join(f"<td>{c}</td>" for c in row) + "</tr>\n"
            html += "</table>"
            tables.append(html)

    return sections, tables
